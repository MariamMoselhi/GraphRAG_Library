"""
DOCX Extractor for HoloLearn
Extracts text content from Word documents.
"""

from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
import json
import time
import re

# Import our utilities
import sys
sys.path.append(str(Path(__file__).parent.parent))
from utils.configs import OUTPUT_DIR, LOGS_DIR
from utils.error_handler import ErrorHandler
from utils.text_cleaner import TextCleaner

try:
    from docx import Document
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


class DOCXExtractor:
    """Extract text from Word documents"""
    
    def __init__(self):
        if not _DOCX_AVAILABLE:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        self.text_cleaner = TextCleaner()
        self.base_output_dir = OUTPUT_DIR
        self.base_logs_dir = LOGS_DIR

        # Ensure base directories exist
        self.base_output_dir.mkdir(parents=True, exist_ok=True)
        self.base_logs_dir.mkdir(parents=True, exist_ok=True)
    
    def _create_resource_name(self, filename: str) -> str:
        """
        Create a clean resource name from filename
        
        Example: "Research Paper - AI Ethics.docx" → "research_paper_ai_ethics"
        """
        # Remove extension
        name = Path(filename).stem
        
        # Convert to lowercase
        name = name.lower()
        
        # Replace spaces and special chars with underscore
        name = re.sub(r'[^\w\s-]', '', name)
        name = re.sub(r'[-\s]+', '_', name)
        
        # Remove leading/trailing underscores
        name = name.strip('_')
        
        # Limit length
        if len(name) > 50:
            name = name[:50]
        
        return name or "unnamed_resource"
    
    def _setup_resource_directories(self, resource_name: str, output_dir_override: Optional[Path] = None) -> tuple:
        """
        Create directories for a specific resource

        Args:
            resource_name: Clean name derived from the input filename.
            output_dir_override: If provided, use this directory for output
                instead of creating a resource-specific subfolder.

        Returns:
            (output_dir, logs_dir) paths
        """
        if output_dir_override:
            resource_output_dir = Path(output_dir_override)
        else:
            resource_output_dir = self.base_output_dir / resource_name

        resource_logs_dir = self.base_logs_dir / resource_name

        resource_output_dir.mkdir(parents=True, exist_ok=True)
        resource_logs_dir.mkdir(parents=True, exist_ok=True)

        return resource_output_dir, resource_logs_dir
    
    def extract(self,
                docx_path: str,
                resource_id: Optional[str] = None,
                clean_text: bool = True,
                include_tables: bool = True,
                preserve_headings: bool = True,
                output_dir: Optional[str] = None) -> Dict[str, Any]:
        """
        Extract text from a Word document

        Args:
            docx_path: Path to DOCX file
            resource_id: Optional unique identifier (if None, uses filename)
            clean_text: Whether to clean the extracted text
            include_tables: Whether to extract text from tables
            preserve_headings: Whether to mark headings in output
            output_dir: Optional shared output directory. When provided,
                all output files are written here instead of a per-resource subfolder.

        Returns:
            Dictionary with extraction results and metadata

        Example:
            extractor = DOCXExtractor()
            result = extractor.extract("notes.docx")
            # Creates: output/notes/notes_text.txt and output/notes/notes_metadata.json
        """
        start_time = time.time()
        docx_path = Path(docx_path)

        # Create resource name from filename
        resource_name = self._create_resource_name(docx_path.name)

        # Setup directories
        override = Path(output_dir) if output_dir else None
        output_dir, logs_dir = self._setup_resource_directories(resource_name, output_dir_override=override)
        
        # Initialize error handler for this specific resource
        error_handler = ErrorHandler(f"docx_{resource_name}")
        # Move log file to resource-specific directory
        error_handler.log_file = logs_dir / "extraction.log"
        error_handler.logger = error_handler._setup_logger()
        
        # Validate file
        if not docx_path.exists():
            error_msg = f"DOCX file not found: {docx_path}"
            error_handler.log_error(
                FileNotFoundError(error_msg),
                context="Validating DOCX file",
                metadata={"path": str(docx_path)}
            )
            return self._create_error_result(
                resource_name, error_msg, output_dir, docx_path.name
            )
        
        # Check file size
        file_size = docx_path.stat().st_size
        file_size_mb = file_size / (1024 * 1024)
        
        error_handler.log_info(
            f"Starting DOCX extraction: {docx_path.name}",
            metadata={
                "size_mb": f"{file_size_mb:.2f}",
                "resource_name": resource_name,
                "output_dir": str(output_dir)
            }
        )
        
        try:
            # Open Word document
            doc = Document(docx_path)
            
            # Extract text from paragraphs
            extracted_text = ""
            paragraph_count = 0
            heading_count = 0
            table_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraph_count += 1
                    
                    # Check if it's a heading
                    if preserve_headings and paragraph.style.name.startswith('Heading'):
                        heading_level = paragraph.style.name.replace('Heading ', '')
                        extracted_text += f"\n{'='*60}\n"
                        extracted_text += f"[HEADING {heading_level}] {paragraph.text}\n"
                        extracted_text += f"{'='*60}\n\n"
                        heading_count += 1
                    else:
                        extracted_text += paragraph.text + "\n"
            
            # Extract text from tables if requested
            if include_tables and doc.tables:
                for table_num, table in enumerate(doc.tables, 1):
                    table_count += 1
                    extracted_text += f"\n--- TABLE {table_num} ---\n"
                    
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        
                        if row_text:
                            extracted_text += " | ".join(row_text) + "\n"
                    
                    extracted_text += "\n"
            
            # Clean text if requested
            if clean_text:
                extracted_text = self.text_cleaner.clean_text(
                    extracted_text,
                    remove_urls=False,  # Keep URLs
                    remove_emails=False,  # Keep emails
                    fix_spacing=True
                )
            
            # Remove duplicate lines
            extracted_text = self.text_cleaner.remove_duplicate_lines(extracted_text)
            
            # Calculate processing time
            processing_time = time.time() - start_time
            
            # Get document properties
            core_props = doc.core_properties
            
            # Create metadata
            metadata = {
                "resource_name": resource_name,
                "resource_id": resource_id or resource_name,
                "filename": docx_path.name,
                "source_type": "docx",
                "upload_date": datetime.now().isoformat(),
                "extraction_timestamp": datetime.now().isoformat(),
                "file_size_bytes": file_size,
                "processing_time_seconds": round(processing_time, 2),
                "status": "success",
                "error_message": None,
                "paragraph_count": paragraph_count,
                "heading_count": heading_count,
                "table_count": table_count,
                "character_count": len(extracted_text),
                "included_tables": include_tables,
                "document_properties": {
                    "title": core_props.title or "N/A",
                    "author": core_props.author or "N/A",
                    "subject": core_props.subject or "N/A",
                    "created": str(core_props.created) if core_props.created else "N/A",
                    "modified": str(core_props.modified) if core_props.modified else "N/A"
                }
            }
            
            # Save text file
            text_file = output_dir / f"{resource_name}_text.txt"
            text_file.write_text(extracted_text, encoding='utf-8')
            metadata["extracted_text_path"] = str(text_file)
            
            # Save metadata file
            metadata_file = output_dir / f"{resource_name}_metadata.json"
            metadata_file.write_text(json.dumps(metadata, indent=2), encoding='utf-8')
            
            error_handler.log_success(
                f"DOCX extracted successfully: {docx_path.name}",
                metadata={
                    "paragraphs": paragraph_count,
                    "headings": heading_count,
                    "tables": table_count,
                    "chars": len(extracted_text),
                    "time": f"{processing_time:.2f}s",
                    "output": str(output_dir)
                }
            )
            
            return {
                "success": True,
                "resource_name": resource_name,
                "resource_id": resource_id or resource_name,
                "text_file": str(text_file),
                "metadata_file": str(metadata_file),
                "output_dir": str(output_dir),
                "logs_dir": str(logs_dir),
                "metadata": metadata,
                "extracted_text": extracted_text
            }
            
        except Exception as e:
            processing_time = time.time() - start_time
            
            error_handler.log_error(
                e,
                context=f"Extracting DOCX: {docx_path.name}",
                metadata={"resource_name": resource_name}
            )
            
            return self._create_error_result(
                resource_name,
                str(e),
                output_dir,
                docx_path.name,
                file_size,
                processing_time
            )
    
    def _create_error_result(self,
                           resource_name: str,
                           error_message: str,
                           output_dir: Path,
                           filename: str = "unknown",
                           file_size: int = 0,
                           processing_time: float = 0) -> Dict[str, Any]:
        """Create error result when extraction fails"""
        
        metadata = {
            "resource_name": resource_name,
            "filename": filename,
            "source_type": "docx",
            "upload_date": datetime.now().isoformat(),
            "extraction_timestamp": datetime.now().isoformat(),
            "file_size_bytes": file_size,
            "processing_time_seconds": round(processing_time, 2),
            "status": "failed",
            "error_message": error_message
        }
        
        # Save metadata even for failed extraction
        metadata_file = output_dir / f"{resource_name}_metadata.json"
        metadata_file.write_text(json.dumps(metadata, indent=2), encoding='utf-8')
        
        return {
            "success": False,
            "resource_name": resource_name,
            "text_file": None,
            "metadata_file": str(metadata_file),
            "output_dir": str(output_dir),
            "metadata": metadata,
            "extracted_text": "",
            "error": error_message
        }
    
    def extract_metadata_only(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract only metadata without extracting text
        Useful for quick file info
        
        Returns:
            Dictionary with DOCX metadata
        """
        try:
            docx_path = Path(docx_path)
            doc = Document(docx_path)
            
            # Count elements
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            table_count = len(doc.tables)
            
            # Get properties
            core_props = doc.core_properties
            
            metadata = {
                "filename": docx_path.name,
                "resource_name": self._create_resource_name(docx_path.name),
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "file_size_bytes": docx_path.stat().st_size,
                "file_size_mb": round(docx_path.stat().st_size / (1024 * 1024), 2),
                "core_properties": {
                    "title": core_props.title or "N/A",
                    "author": core_props.author or "N/A",
                    "subject": core_props.subject or "N/A",
                    "created": str(core_props.created) if core_props.created else "N/A",
                    "modified": str(core_props.modified) if core_props.modified else "N/A"
                }
            }
            
            return metadata
            
        except Exception as e:
            error_handler = ErrorHandler("docx_metadata")
            error_handler.log_error(
                e,
                context=f"Extracting metadata from {docx_path}"
            )
            return {}


# Example usage and testing
if __name__ == "__main__":
    from utils.file_picker import FilePicker
    
    print("=== Testing DOCX Extractor ===\n")
    
    # Initialize extractor
    extractor = DOCXExtractor()
    
    # Use file picker to select DOCX
    picker = FilePicker()
    print("Please select a Word document...")
    test_docx = picker.pick_docx()
    picker.close()
    
    if test_docx:
        print(f"\n✓ Selected: {Path(test_docx).name}\n")
        
        print("1. Extracting metadata only...")
        metadata = extractor.extract_metadata_only(test_docx)
        print(f"   Resource name: {metadata.get('resource_name', 'N/A')}")
        print(f"   Paragraphs: {metadata.get('paragraph_count', 'N/A')}")
        print(f"   Tables: {metadata.get('table_count', 'N/A')}")
        print(f"   Size: {metadata.get('file_size_mb', 'N/A')} MB")
        print(f"   Title: {metadata.get('core_properties', {}).get('title', 'N/A')}")
        print(f"   Author: {metadata.get('core_properties', {}).get('author', 'N/A')}\n")
        
        print("2. Full extraction (with tables and headings)...")
        result = extractor.extract(
            docx_path=test_docx,
            clean_text=True,
            include_tables=True,
            preserve_headings=True
        )
        
        if result['success']:
            print(f"   ✓ Success!")
            print(f"   Resource name: {result['resource_name']}")
            print(f"   Output directory: {result['output_dir']}")
            print(f"   Logs directory: {result['logs_dir']}")
            print(f"   Text file: {result['text_file']}")
            print(f"   Metadata file: {result['metadata_file']}")
            print(f"   Paragraphs: {result['metadata']['paragraph_count']}")
            print(f"   Headings: {result['metadata']['heading_count']}")
            print(f"   Tables: {result['metadata']['table_count']}")
            print(f"   Characters: {result['metadata']['character_count']}")
            print(f"   Processing time: {result['metadata']['processing_time_seconds']}s")
            
            print(f"\n   Preview (first 400 chars):")
            print("   " + "-" * 50)
            preview = result['extracted_text'][:400]
            print(f"   {preview}...")
            print("   " + "-" * 50)
        else:
            print(f"   ✗ Failed: {result['error']}")
    else:
        print("❌ No file selected")
        print("   The extractor is ready to use!")
# ```

# ---

# ## **Features:**

# ✅ **Extracts paragraphs** - All text content from document  
# ✅ **Preserves headings** - Marks heading levels (H1, H2, etc.)  
# ✅ **Table extraction** - Converts tables to text format  
# ✅ **Document properties** - Title, author, dates  
# ✅ **Resource-based folders** - Same naming as PDF/PPTX  
# ✅ **File picker support** - Easy testing  

# ---

# ## **Output Example:**

# **Input:** `"Study Notes - Chapter 3.docx"`

# **Creates:**
# ```
# output/study_notes_chapter_3/
#   study_notes_chapter_3_text.txt
#   study_notes_chapter_3_metadata.json
  
# logs/study_notes_chapter_3/
#   extraction.log
# ```

# **study_notes_chapter_3_text.txt:**
# ```
# ============================================================
# [HEADING 1] Introduction to Neural Networks
# ============================================================

# Neural networks are computational models inspired by biological neurons...

# ============================================================
# [HEADING 2] Key Concepts
# ============================================================

# The main components include:
# - Input layer
# - Hidden layers
# - Output layer

# --- TABLE 1 ---
# Layer Type | Neurons | Activation Function
# Input | 784 | None
# Hidden 1 | 128 | ReLU
# Hidden 2 | 64 | ReLU
# Output | 10 | Softmax